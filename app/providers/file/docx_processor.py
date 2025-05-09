from typing import BinaryIO, Dict, Any
from docx import Document
from app.providers.file.base_file_processor import FileContentProcessor
from app.utils.helpers import extract_topics

class DOCXProcessor(FileContentProcessor):
    """Processor for DOCX files"""
    
    def process_file(self, file: BinaryIO) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        # Handle FastAPI UploadFile objects
        if hasattr(file, 'file'):
            file_obj = file.file
            # Store filename for later use
            filename = getattr(file, 'filename', 'unknown.docx')
        else:
            file_obj = file
            filename = getattr(file, 'filename', 'unknown.docx')
            
        doc = Document(file_obj)
        
        # Extract text from paragraphs
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        text = "\n".join(paragraphs)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        # Extract metadata from document properties
        metadata = {}
        props = doc.core_properties
        
        if props:
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "last_modified_by": props.last_modified_by or ""
            }
        

        return {
            "content": text,
            "metadata": {
                **metadata,
                "source": "file_upload",
                "file_type": "docx",
                "extraction_method": "python-docx",
                "filename": filename
            }
        }